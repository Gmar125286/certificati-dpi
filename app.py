from __future__ import annotations

import hashlib
import hmac
import json
import os
import re
import secrets
import shutil
import sqlite3
import subprocess
import sys
import threading
import traceback
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin
from urllib.request import Request, urlopen

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def application_base_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


BASE_DIR = application_base_dir()
TEMPLATES_DIR = BASE_DIR
OUTPUT_DIR = BASE_DIR / "output"
DOCX_OUTPUT_DIR = OUTPUT_DIR / "docx"
PDF_OUTPUT_DIR = OUTPUT_DIR / "pdf"
ASSETS_OUTPUT_DIR = OUTPUT_DIR / "assets"
LOG_PATH = OUTPUT_DIR / "error.log"
DB_PATH = BASE_DIR / "storico_revisioni.db"
SETTINGS_PATH = BASE_DIR / "settings.json"
USERS_PATH = BASE_DIR / "users.json"
IRUDEK_NORMS_PATH = BASE_DIR / "irudek_norme.json"
ICON_PATH = BASE_DIR / "gestione_certificati_dpi.ico"
DEFAULT_LOGO_CANDIDATES = [
    BASE_DIR / "logo con scrittura .jpg",
    BASE_DIR / "logo_azienda.jpg",
]
PASSWORD_ITERATIONS = 200000
DEFAULT_ADMIN_USERNAME = "admin"
DEFAULT_ADMIN_SALT = "6baf5c2a4f7e9d31b8c6e12a5d9f0b44"
DEFAULT_ADMIN_HASH = "31904316b7ff181c030bdc224d9743e1b4de8c3aad0a4876075ea2eb31bcc402"
IRUDEK_BASE_URL = "https://www.irudek.com"
IRUDEK_CATALOG_SEEDS = [
    "/en/cat/fall-arrest-protection/",
    "/en/cat/lifelines/",
]
IRUDEK_CATALOG_MAX_AGE_DAYS = 7
STAMP_TEXT_LINES = [
    "ZENIT SRL UNIPERSONALE",
    "Via del Tratturello Tarantino, 5",
    "74123 TARANTO",
    "Tel. 099/4725984 - Fax 099/4723444",
    "Partita IVA 02455090734",
]


def slugify(value: str) -> str:
    cleaned = "".join(ch if ch.isalnum() else "_" for ch in value.strip().lower())
    while "__" in cleaned:
        cleaned = cleaned.replace("__", "_")
    return cleaned.strip("_") or "documento"


def normalize_label(value: str) -> str:
    return " ".join(value.replace("\n", " ").strip().upper().split())


def template_files() -> list[Path]:
    return sorted(TEMPLATES_DIR.glob("*.docx"), key=lambda path: path.name.lower())


def extract_checklist_items(template_path: Path) -> list[str]:
    document = Document(template_path)
    if not document.tables:
        return []

    for table in reversed(document.tables):
        if len(table.columns) < 4 or len(table.rows) < 4:
            continue
        header = [normalize_label(cell.text) for cell in table.rows[0].cells]
        if header and header[0] == "MODELLO D.P.I.":
            items = []
            for row in table.rows[2:]:
                if len(row.cells) < 4:
                    continue
                description = row.cells[1].text.replace("\n", " ").strip()
                if description:
                    items.append(description)
            return items
    return []


def default_settings() -> dict:
    for candidate in DEFAULT_LOGO_CANDIDATES:
        if candidate.exists():
            return {"logo_path": str(candidate)}
    return {"logo_path": ""}


def current_datetime_display() -> str:
    return datetime.now().strftime("%d/%m/%Y %H:%M")


def date_only_display(value: str) -> str:
    return value.strip().split(" ")[0] if value.strip() else ""


def parse_catalog_datetime(value: str) -> datetime | None:
    try:
        return datetime.strptime(value.strip(), "%Y-%m-%d %H:%M:%S")
    except Exception:
        return None


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def hidden_subprocess_kwargs() -> dict:
    if os.name != "nt":
        return {}
    startupinfo = subprocess.STARTUPINFO()
    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
    return {
        "startupinfo": startupinfo,
        "creationflags": subprocess.CREATE_NO_WINDOW,
    }


def _load_font(candidates: list[Path], size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def normalize_username(username: str) -> str:
    return username.strip().lower()


def normalize_catalog_code(value: str) -> str:
    return "".join(str(value).strip().upper().split())


def hash_password(password: str, salt_hex: str, iterations: int = PASSWORD_ITERATIONS) -> str:
    return hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        bytes.fromhex(salt_hex),
        iterations,
    ).hex()


def fetch_web_text(url: str) -> str:
    request = Request(url, headers={"User-Agent": "Mozilla/5.0"})
    with urlopen(request, timeout=25) as response:
        return response.read().decode("utf-8", errors="ignore")


def extract_irudek_products_array(html: str) -> list[dict]:
    marker = '"products":['
    index = html.find(marker)
    if index == -1:
        return []
    start = html.find("[", index)
    nesting = 0
    for pos, char in enumerate(html[start:], start):
        if char == "[":
            nesting += 1
        elif char == "]":
            nesting -= 1
            if nesting == 0:
                return json.loads(html[start : pos + 1])
    return []


def split_irudek_refs(reference_value: str) -> list[str]:
    refs: list[str] = []
    for part in re.split(r"[,;/]+", str(reference_value or "")):
        cleaned = normalize_catalog_code(part)
        if cleaned:
            refs.append(cleaned)
    return refs


def build_irudek_norm_catalog(output_path: Path) -> tuple[int, int]:
    category_paths = set(IRUDEK_CATALOG_SEEDS)
    for seed in IRUDEK_CATALOG_SEEDS:
        html = fetch_web_text(urljoin(IRUDEK_BASE_URL, seed))
        for link in re.findall(r'"(/en/cat/[^"]+/?)"', html):
            if link.startswith(seed) and link != seed:
                category_paths.add(link)

    entries: dict[str, dict] = {}
    conflicts = 0

    for path in sorted(category_paths):
        try:
            html = fetch_web_text(urljoin(IRUDEK_BASE_URL, path))
        except Exception:
            continue
        for product in extract_irudek_products_array(html):
            title = str(product.get("title", "")).strip()
            slug = str(product.get("slug", "")).strip()
            general = product.get("general", {}) or {}
            standard = ""
            specs = general.get("specs", {}) or {}
            for item in specs.get("items", []) or []:
                if str(item.get("key", "")).strip().lower() == "standard":
                    standard = str(item.get("value", "")).strip()
                    break
            if not standard:
                continue
            for ref in split_irudek_refs(general.get("reference", "")):
                existing = entries.get(ref)
                if existing and existing.get("standard") != standard:
                    conflicts += 1
                    continue
                entries[ref] = {
                    "standard": standard,
                    "title": title,
                    "slug": slug,
                }

    payload = {
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "source": [urljoin(IRUDEK_BASE_URL, path) for path in sorted(category_paths)],
        "total_codes": len(entries),
        "items": dict(sorted(entries.items())),
    }
    output_path.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    return len(entries), conflicts


def generate_stamp_signature_asset(reviewer_name: str) -> Path:
    ASSETS_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    canvas_width = 1200
    canvas_height = 620
    image = Image.new("RGBA", (canvas_width, canvas_height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)

    bold_font = _load_font(
        [Path(r"C:\Windows\Fonts\arialbd.ttf"), Path(r"C:\Windows\Fonts\Arialbd.ttf")],
        54,
    )
    regular_font = _load_font(
        [Path(r"C:\Windows\Fonts\arial.ttf"), Path(r"C:\Windows\Fonts\Arial.ttf")],
        38,
    )
    script_font = _load_font(
        [
            Path(r"C:\Windows\Fonts\segoesc.ttf"),
            Path(r"C:\Windows\Fonts\BRUSHSCI.TTF"),
            Path(r"C:\Windows\Fonts\SCRIPTBL.TTF"),
        ],
        110,
    )

    y = 20
    x_center = canvas_width // 2
    ink = (20, 20, 20, 245)

    title_bbox = draw.textbbox((0, 0), STAMP_TEXT_LINES[0], font=bold_font)
    title_width = title_bbox[2] - title_bbox[0]
    draw.text(((canvas_width - title_width) / 2, y), STAMP_TEXT_LINES[0], fill=ink, font=bold_font)
    y += 92

    for line in STAMP_TEXT_LINES[1:]:
        bbox = draw.textbbox((0, 0), line, font=regular_font)
        width = bbox[2] - bbox[0]
        draw.text(((canvas_width - width) / 2, y), line, fill=ink, font=regular_font)
        y += 58

    reviewer_clean = " ".join(part.capitalize() for part in reviewer_name.strip().split()) or "Revisore"
    signature_text = reviewer_clean
    signature_layer = Image.new("RGBA", (canvas_width, canvas_height), (255, 255, 255, 0))
    signature_draw = ImageDraw.Draw(signature_layer)
    sig_fill = (12, 12, 12, 210)
    sig_bbox = signature_draw.textbbox((0, 0), signature_text, font=script_font)
    sig_w = sig_bbox[2] - sig_bbox[0]
    sig_h = sig_bbox[3] - sig_bbox[1]
    sig_x = int(x_center - sig_w / 2)
    sig_y = int((canvas_height - sig_h) / 2) - 10

    for dx, dy in [(0, 0), (1, 0), (0, 1)]:
        signature_draw.text((sig_x + dx, sig_y + dy), signature_text, fill=sig_fill, font=script_font)

    signature_layer = signature_layer.rotate(-8, resample=Image.Resampling.BICUBIC, expand=False)
    image = Image.alpha_composite(image, signature_layer)

    content_bbox = image.getbbox()
    if content_bbox:
        image = image.crop(content_bbox)

    output_path = ASSETS_OUTPUT_DIR / f"timbro_firma_{slugify(reviewer_name)}.png"
    image.save(output_path)
    return output_path


class UserStore:
    def __init__(self, file_path: Path) -> None:
        self.file_path = file_path
        self._ensure_store()

    def _ensure_store(self) -> None:
        if self.file_path.exists():
            try:
                payload = json.loads(self.file_path.read_text(encoding="utf-8"))
                if isinstance(payload.get("users"), list):
                    return
            except Exception:
                pass

        self._write_users(
            [
                {
                    "username": DEFAULT_ADMIN_USERNAME,
                    "salt": DEFAULT_ADMIN_SALT,
                    "password_hash": DEFAULT_ADMIN_HASH,
                    "iterations": PASSWORD_ITERATIONS,
                    "created_at": datetime.now().isoformat(timespec="seconds"),
                }
            ]
        )

    def _read_users(self) -> list[dict]:
        try:
            payload = json.loads(self.file_path.read_text(encoding="utf-8"))
        except Exception:
            self._ensure_store()
            payload = json.loads(self.file_path.read_text(encoding="utf-8"))
        users = payload.get("users", [])
        return users if isinstance(users, list) else []

    def _write_users(self, users: list[dict]) -> None:
        self.file_path.write_text(json.dumps({"users": users}, indent=2), encoding="utf-8")

    def verify_credentials(self, username: str, password: str) -> bool:
        normalized = normalize_username(username)
        for user in self._read_users():
            if normalize_username(str(user.get("username", ""))) != normalized:
                continue
            try:
                actual_hash = hash_password(
                    password,
                    str(user.get("salt", "")),
                    int(user.get("iterations", PASSWORD_ITERATIONS)),
                )
            except Exception:
                return False
            expected_hash = str(user.get("password_hash", ""))
            return hmac.compare_digest(actual_hash, expected_hash)
        return False

    def create_user(self, username: str, password: str) -> str:
        cleaned_username = username.strip()
        normalized = normalize_username(cleaned_username)
        if not normalized:
            raise ValueError("Inserisci un nome utente.")
        if len(password) < 4:
            raise ValueError("La password deve contenere almeno 4 caratteri.")

        users = self._read_users()
        if any(normalize_username(str(user.get("username", ""))) == normalized for user in users):
            raise ValueError("Esiste gia' un utente con questo nome.")

        salt = secrets.token_hex(16)
        users.append(
            {
                "username": cleaned_username,
                "salt": salt,
                "password_hash": hash_password(password, salt),
                "iterations": PASSWORD_ITERATIONS,
                "created_at": datetime.now().isoformat(timespec="seconds"),
            }
        )
        users.sort(key=lambda item: normalize_username(str(item.get("username", ""))))
        self._write_users(users)
        return cleaned_username

    def list_usernames(self) -> list[str]:
        return [str(user.get("username", "")).strip() for user in self._read_users() if str(user.get("username", "")).strip()]


class IrudekNormCatalog:
    def __init__(self, file_path: Path) -> None:
        self.file_path = file_path
        self.items: dict[str, dict] = {}
        self.generated_at = ""
        self.load()

    def load(self) -> None:
        self.items = {}
        self.generated_at = ""
        if not self.file_path.exists():
            return
        try:
            payload = json.loads(self.file_path.read_text(encoding="utf-8"))
        except Exception:
            return
        self.generated_at = str(payload.get("generated_at", "")).strip()
        raw_items = payload.get("items", {})
        if not isinstance(raw_items, dict):
            return
        for code, entry in raw_items.items():
            normalized = normalize_catalog_code(code)
            if not normalized or not isinstance(entry, dict):
                continue
            self.items[normalized] = entry

    def lookup(self, code: str) -> dict | None:
        return self.items.get(normalize_catalog_code(code))

    def search_by_name(self, query: str, limit: int = 20) -> list[tuple[str, dict]]:
        cleaned = " ".join(query.strip().upper().split())
        if not cleaned:
            return []
        terms = [term for term in cleaned.split() if term]
        matches: list[tuple[int, tuple[str, dict]]] = []
        for code, entry in self.items.items():
            title = str(entry.get("title", "")).upper()
            title_terms = title.split()
            if not all(any(token.startswith(term) for token in title_terms) for term in terms):
                continue
            score = 0
            if title == cleaned:
                score += 200
            if title.startswith(cleaned):
                score += 120
            if cleaned in title:
                score += 60
            score += max(0, 30 - abs(len(title) - len(cleaned)))
            matches.append((score, (code, entry)))
        matches.sort(key=lambda item: (-item[0], item[1][1].get("title", ""), item[1][0]))
        return [item for _, item in matches[:limit]]

    def is_stale(self, max_age_days: int = IRUDEK_CATALOG_MAX_AGE_DAYS) -> bool:
        generated = parse_catalog_datetime(self.generated_at)
        if generated is None:
            return True
        return (datetime.now() - generated).days >= max_age_days


@dataclass
class RevisionRecord:
    cliente: str
    template_name: str
    modello_dispositivo: str
    norma: str
    codice_sap: str
    lotto: str
    serie: str
    anno_fabbricazione: str
    data_revisione: str
    data_prossima_revisione: str
    tecnico: str
    osservazioni: str
    esito: str


class RevisionRepository:
    def __init__(self, db_path: Path) -> None:
        self.db_path = db_path
        self._ensure_db()

    def _connect(self) -> sqlite3.Connection:
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn

    def _ensure_db(self) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS revisioni (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    created_at TEXT NOT NULL,
                    updated_at TEXT,
                    cliente TEXT NOT NULL,
                    template_name TEXT NOT NULL,
                    modello_dispositivo TEXT,
                    norma TEXT,
                    codice_sap TEXT,
                    lotto TEXT,
                    serie TEXT,
                    anno_fabbricazione TEXT,
                    data_revisione TEXT,
                    data_prossima_revisione TEXT,
                    tecnico TEXT,
                    osservazioni TEXT,
                    esito TEXT,
                    docx_path TEXT,
                    pdf_path TEXT
                )
                """
            )
            columns = {row["name"] for row in conn.execute("PRAGMA table_info(revisioni)").fetchall()}
            if "updated_at" not in columns:
                conn.execute("ALTER TABLE revisioni ADD COLUMN updated_at TEXT")

    def insert_record(
        self,
        record: RevisionRecord,
        docx_path: Path | None,
        pdf_path: Path | None,
    ) -> int:
        with self._connect() as conn:
            now_value = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            cursor = conn.execute(
                """
                INSERT INTO revisioni (
                    created_at,
                    updated_at,
                    cliente,
                    template_name,
                    modello_dispositivo,
                    norma,
                    codice_sap,
                    lotto,
                    serie,
                    anno_fabbricazione,
                    data_revisione,
                    data_prossima_revisione,
                    tecnico,
                    osservazioni,
                    esito,
                    docx_path,
                    pdf_path
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    now_value,
                    now_value,
                    record.cliente,
                    record.template_name,
                    record.modello_dispositivo,
                    record.norma,
                    record.codice_sap,
                    record.lotto,
                    record.serie,
                    record.anno_fabbricazione,
                    record.data_revisione,
                    record.data_prossima_revisione,
                    record.tecnico,
                    record.osservazioni,
                    record.esito,
                    str(docx_path) if docx_path else "",
                    str(pdf_path) if pdf_path else "",
                ),
            )
            return int(cursor.lastrowid)

    def search_records(self, query: str = "") -> list[sqlite3.Row]:
        cleaned = query.strip().lower()
        with self._connect() as conn:
            if not cleaned:
                return list(
                    conn.execute(
                        """
                        SELECT id, created_at, updated_at, cliente, modello_dispositivo, serie, data_revisione,
                               esito, template_name, docx_path, pdf_path
                        FROM revisioni
                        ORDER BY COALESCE(updated_at, created_at) DESC, id DESC
                        """
                    )
                )

            terms = [part.strip() for part in cleaned.split("&&") if part.strip()]
            if not terms:
                terms = [cleaned]

            searchable_fields = [
                "LOWER(cliente)",
                "LOWER(COALESCE(lotto, ''))",
                "LOWER(COALESCE(serie, ''))",
                "LOWER(modello_dispositivo)",
                "LOWER(template_name)",
            ]
            clause_template = "(" + " OR ".join(f"{field} LIKE ?" for field in searchable_fields) + ")"
            where_clause = " AND ".join(clause_template for _ in terms)
            parameters: list[str] = []
            for term in terms:
                like_value = f"%{term}%"
                parameters.extend([like_value] * len(searchable_fields))

            return list(
                conn.execute(
                    f"""
                    SELECT id, created_at, updated_at, cliente, modello_dispositivo, serie, data_revisione,
                           esito, template_name, docx_path, pdf_path
                    FROM revisioni
                    WHERE {where_clause}
                    ORDER BY COALESCE(updated_at, created_at) DESC, id DESC
                    """,
                    parameters,
                )
            )

    def distinct_clients(self) -> list[str]:
        with self._connect() as conn:
            rows = conn.execute(
                """
                SELECT DISTINCT cliente
                FROM revisioni
                WHERE TRIM(COALESCE(cliente, '')) <> ''
                ORDER BY cliente COLLATE NOCASE
                """
            ).fetchall()
        return [row["cliente"] for row in rows]

    def latest_record_for_client(self, cliente: str) -> sqlite3.Row | None:
        cleaned = cliente.strip()
        if not cleaned:
            return None
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT *
                FROM revisioni
                WHERE LOWER(cliente) = LOWER(?)
                ORDER BY COALESCE(updated_at, created_at) DESC, id DESC
                LIMIT 1
                """,
                (cleaned,),
            ).fetchone()

    def get_record_by_id(self, record_id: int) -> sqlite3.Row | None:
        with self._connect() as conn:
            return conn.execute(
                """
                SELECT *
                FROM revisioni
                WHERE id = ?
                """,
                (record_id,),
            ).fetchone()

    def delete_record(self, record_id: int) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                DELETE FROM revisioni
                WHERE id = ?
                """,
                (record_id,),
            )

    def update_record(
        self,
        record_id: int,
        record: RevisionRecord,
        docx_path: Path | None,
        pdf_path: Path | None,
    ) -> None:
        with self._connect() as conn:
            conn.execute(
                """
                UPDATE revisioni
                SET updated_at = ?,
                    cliente = ?,
                    template_name = ?,
                    modello_dispositivo = ?,
                    norma = ?,
                    codice_sap = ?,
                    lotto = ?,
                    serie = ?,
                    anno_fabbricazione = ?,
                    data_revisione = ?,
                    data_prossima_revisione = ?,
                    tecnico = ?,
                    osservazioni = ?,
                    esito = ?,
                    docx_path = ?,
                    pdf_path = ?
                WHERE id = ?
                """,
                (
                    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    record.cliente,
                    record.template_name,
                    record.modello_dispositivo,
                    record.norma,
                    record.codice_sap,
                    record.lotto,
                    record.serie,
                    record.anno_fabbricazione,
                    record.data_revisione,
                    record.data_prossima_revisione,
                    record.tecnico,
                    record.osservazioni,
                    record.esito,
                    str(docx_path) if docx_path else "",
                    str(pdf_path) if pdf_path else "",
                    record_id,
                ),
            )


class CertificateGenerator:
    FIELD_LABELS = {
        "CLIENTE": "cliente",
        "MODELLO DISPOSITIVO": "modello_dispositivo",
        "NORMA": "norma",
        "CODICE ARTICOLO SAP": "codice_sap",
        "NUMERO DI LOTTO": "lotto",
        "NUMERO DI SERIE": "serie",
        "ANNO DI FABBRICAZIONE": "anno_fabbricazione",
    }

    def generate(
        self,
        template_path: Path,
        record: RevisionRecord,
        export_pdf: bool,
        logo_path: Path | None = None,
        checklist_results: list[str] | None = None,
    ) -> tuple[Path, Path | None]:
        DOCX_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        PDF_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_base = (
            f"{timestamp}_{slugify(record.cliente)}_{slugify(record.modello_dispositivo or template_path.stem)}"
        )
        docx_path = DOCX_OUTPUT_DIR / f"{filename_base}.docx"
        shutil.copy2(template_path, docx_path)

        document = Document(docx_path)
        self._adjust_page_layout(document)
        self._insert_logo(document, logo_path)
        self._fill_tables(document, record, checklist_results or [])
        self._fill_signature_line(document, record)
        document.save(docx_path)

        pdf_path = None
        pdf_error = None
        if export_pdf:
            pdf_path = PDF_OUTPUT_DIR / f"{filename_base}.pdf"
            try:
                self._convert_to_pdf(docx_path, pdf_path)
            except Exception as exc:
                pdf_error = str(exc)
                pdf_path = None

        return docx_path, pdf_path, pdf_error

    def print_document(self, docx_path: Path, copies: int = 2) -> None:
        escaped_docx = str(docx_path).replace("'", "''")
        copies = max(1, int(copies))
        command = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
    $document = $word.Documents.Open('{escaped_docx}')
    for ($i = 0; $i -lt {copies}; $i++) {{
        $document.PrintOut()
    }}
    $document.Close($false)
}} finally {{
    $word.Quit()
}}
"""
        completed = subprocess.run(
            ["powershell", "-NoProfile", "-Command", command],
            capture_output=True,
            text=True,
            check=False,
            **hidden_subprocess_kwargs(),
        )
        if completed.returncode != 0:
            raise RuntimeError(
                "Stampa non riuscita. "
                + (completed.stderr.strip() or completed.stdout.strip() or "Errore sconosciuto.")
            )

    def _adjust_page_layout(self, document: Document) -> None:
        for section in document.sections:
            section.top_margin = Cm(1.2)
            section.bottom_margin = Cm(1.6)

    def _insert_logo(self, document: Document, logo_path: Path | None) -> None:
        if not logo_path or not logo_path.exists():
            return

        anchor = document.paragraphs[0] if document.paragraphs else document.add_paragraph()
        logo_paragraph = anchor.insert_paragraph_before()
        logo_paragraph.alignment = 0
        logo_paragraph.paragraph_format.space_before = 0
        logo_paragraph.paragraph_format.space_after = 0
        logo_paragraph.add_run().add_picture(str(logo_path), width=Cm(6.5))

    def _fill_tables(self, document: Document, record: RevisionRecord, checklist_results: list[str]) -> None:
        for table in document.tables:
            self._fill_revision_table(table, record)
            self._fill_device_info_table(table, record)
            self._fill_esito_row(table, record.esito)
            self._fill_checklist_table(table, checklist_results, record.esito)

    def _write_cell_value(self, cell, value: str, *, bold: bool = True, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align
        run = paragraph.add_run(value)
        run.bold = bold

    def _fill_revision_table(self, table, record: RevisionRecord) -> None:
        cells = [[normalize_label(cell.text) for cell in row.cells] for row in table.rows]
        if not cells:
            return

        if len(cells) >= 4 and len(cells[0]) >= 3:
            if "DATA REVISIONE" in cells[1] and "DATA PROSSIMA REVISIONE" in cells[1]:
                row = table.rows[2]
                self._write_cell_value(row.cells[0], date_only_display(record.data_revisione))
                self._write_cell_value(row.cells[1], record.data_prossima_revisione)
                self._write_cell_value(row.cells[2], record.tecnico)
                for cell in table.rows[3].cells:
                    if normalize_label(cell.text).startswith("OSSERVAZIONI:"):
                        self._write_cell_value(
                            cell,
                            f"OSSERVAZIONI: {record.osservazioni}".strip(),
                            align=WD_ALIGN_PARAGRAPH.LEFT,
                        )
                return

        if len(cells) >= 3 and len(cells[0]) >= 2:
            if cells[0][0] == "DATA REVISIONE" and cells[0][1] == "TECNICO":
                self._write_cell_value(table.rows[1].cells[0], date_only_display(record.data_revisione))
                self._write_cell_value(table.rows[1].cells[1], record.tecnico)
                for cell in table.rows[2].cells:
                    if normalize_label(cell.text).startswith("OSSERVAZIONI:"):
                        self._write_cell_value(
                            cell,
                            f"OSSERVAZIONI: {record.osservazioni}".strip(),
                            align=WD_ALIGN_PARAGRAPH.LEFT,
                        )

    def _fill_device_info_table(self, table, record: RevisionRecord) -> None:
        record_map = {
            "cliente": record.cliente,
            "modello_dispositivo": record.modello_dispositivo,
            "norma": record.norma,
            "codice_sap": record.codice_sap,
            "lotto": record.lotto,
            "serie": record.serie,
            "anno_fabbricazione": record.anno_fabbricazione,
        }

        for row in table.rows:
            if not row.cells:
                continue
            label = normalize_label(row.cells[0].text)
            field_name = self.FIELD_LABELS.get(label)
            if not field_name:
                continue
            value = record_map[field_name]
            for cell in row.cells[1:]:
                self._write_cell_value(cell, value)

    def _fill_esito_row(self, table, esito: str) -> None:
        mark_positive = esito == "Positivo"
        for row in table.rows:
            if not row.cells:
                continue
            label = normalize_label(row.cells[0].text)
            if label != "ESITO DELLA REVISIONE":
                continue
            if len(row.cells) >= 3 and (
                normalize_label(row.cells[1].text) == "" or normalize_label(row.cells[2].text) == ""
            ):
                self._write_cell_value(row.cells[1], "X" if mark_positive else "")
                self._write_cell_value(row.cells[2], "" if mark_positive else "X")

    def _fill_checklist_table(self, table, checklist_results: list[str], default_esito: str) -> None:
        if len(table.columns) < 4 or len(table.rows) < 4:
            return
        header = [normalize_label(cell.text) for cell in table.rows[0].cells]
        if not header or header[0] != "MODELLO D.P.I.":
            return

        default_value = "Positivo" if default_esito == "Positivo" else "Negativo"
        result_index = 0
        for row in table.rows[2:]:
            if len(row.cells) < 4:
                continue
            description = row.cells[1].text.replace("\n", " ").strip()
            if not description:
                continue
            value = checklist_results[result_index] if result_index < len(checklist_results) else default_value
            self._write_cell_value(row.cells[2], "X" if value == "Positivo" else "")
            self._write_cell_value(row.cells[3], "" if value == "Positivo" else "X")
            result_index += 1

    def _fill_signature_line(self, document: Document, record: RevisionRecord) -> None:
        for paragraph in document.paragraphs:
            text = normalize_label(paragraph.text)
            if "DATA" in text and "IL REVISORE" in text:
                paragraph.text = ""
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                section = document.sections[0]
                usable_width = section.page_width - section.left_margin - section.right_margin
                tab_stops = paragraph.paragraph_format.tab_stops
                tab_stops.clear_all()
                tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

                date_run = paragraph.add_run(f"DATA {record.data_revisione}")
                date_run.bold = True
                paragraph.add_run("\t")
                reviewer_run = paragraph.add_run(f"IL REVISORE {record.tecnico}")
                reviewer_run.bold = True
                return

    def _convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        escaped_docx = str(docx_path).replace("'", "''")
        escaped_pdf = str(pdf_path).replace("'", "''")
        command = f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
$word.DisplayAlerts = 0
try {{
    if (Test-Path '{escaped_pdf}') {{
        Remove-Item -LiteralPath '{escaped_pdf}' -Force
    }}
    $document = $word.Documents.Open('{escaped_docx}', $false, $false)
    try {{
        $document.ExportAsFixedFormat('{escaped_pdf}', 17)
    }} finally {{
        $document.Close($false)
    }}
    if (-not (Test-Path '{escaped_pdf}')) {{
        throw 'PDF non creato da Word.'
    }}
}} catch {{
    Write-Output ('PDF_EXPORT_ERROR: ' + $_.Exception.Message)
    exit 1
}} finally {{
    $word.Quit()
}}
"""
        last_error = "Errore sconosciuto."
        for attempt in range(2):
            completed = subprocess.run(
                ["powershell", "-NoProfile", "-Command", command],
                capture_output=True,
                text=True,
                check=False,
                **hidden_subprocess_kwargs(),
            )
            if completed.returncode == 0 and pdf_path.exists():
                return
            last_error = completed.stderr.strip() or completed.stdout.strip() or "Errore sconosciuto."

        raise RuntimeError(f"Esportazione PDF non riuscita. {last_error}")


class RevisionApp(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Gestione Certificati DPI")
        self.minsize(1120, 720)
        self.is_authenticated = False
        self.current_user = ""
        self.editing_record_id: int | None = None

        self.repository = RevisionRepository(DB_PATH)
        self.user_store = UserStore(USERS_PATH)
        self.norm_catalog = IrudekNormCatalog(IRUDEK_NORMS_PATH)
        self.generator = CertificateGenerator()
        self.templates = template_files()
        self.template_lookup = {path.name: path for path in self.templates}
        self.settings = self._load_settings()

        self.client_widget: ttk.Combobox | None = None
        self.osservazioni_widget: tk.Text | None = None
        self.checklist_items: list[str] = []
        self.checklist_vars: list[tk.StringVar] = []
        self.generate_button: ttk.Button | None = None
        self.generate_print_button: ttk.Button | None = None
        self.catalog_button: ttk.Button | None = None

        self._build_variables()
        self._build_layout()
        self._refresh_client_values()
        self._refresh_history()
        self._sync_model_from_template()
        self._apply_window_icon()
        self._refresh_catalog_info()
        self.attributes("-alpha", 0.0)
        self.update_idletasks()
        if self._show_login_dialog():
            self.is_authenticated = True
            self.attributes("-alpha", 1.0)
            self._maximize_on_start()
        else:
            self.after(0, self.destroy)

    def _maximize_on_start(self) -> None:
        try:
            self.state("zoomed")
        except tk.TclError:
            self.attributes("-zoomed", True)

    def _show_login_dialog(self) -> bool:
        dialog = tk.Toplevel(self)
        dialog.title("Accesso")
        dialog.resizable(False, False)
        dialog.transient(self)
        dialog.protocol("WM_DELETE_WINDOW", dialog.destroy)

        if ICON_PATH.exists():
            try:
                dialog.iconbitmap(default=str(ICON_PATH))
            except Exception:
                pass

        frame = ttk.Frame(dialog, padding=18)
        frame.grid(row=0, column=0, sticky="nsew")
        frame.columnconfigure(1, weight=1)

        ttk.Label(frame, text="Accesso applicazione", font=("Segoe UI", 12, "bold")).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(0, 12)
        )
        ttk.Label(frame, text="Utente").grid(row=1, column=0, sticky="w", padx=(0, 10), pady=6)
        ttk.Label(frame, text="Password").grid(row=2, column=0, sticky="w", padx=(0, 10), pady=6)

        username_var = tk.StringVar()
        password_var = tk.StringVar()
        error_var = tk.StringVar()
        authenticated = {"value": False}

        username_entry = ttk.Entry(frame, textvariable=username_var, width=28)
        password_entry = ttk.Entry(frame, textvariable=password_var, show="*", width=28)
        username_entry.grid(row=1, column=1, sticky="ew", pady=6)
        password_entry.grid(row=2, column=1, sticky="ew", pady=6)
        ttk.Label(frame, textvariable=error_var, foreground="#b42318").grid(
            row=3, column=0, columnspan=2, sticky="w", pady=(4, 8)
        )

        button_row = ttk.Frame(frame)
        button_row.grid(row=4, column=0, columnspan=2, sticky="e", pady=(6, 0))

        def attempt_login(event: object | None = None) -> None:
            if self.user_store.verify_credentials(username_var.get(), password_var.get()):
                authenticated["value"] = True
                self.current_user = username_var.get().strip()
                dialog.destroy()
                return
            error_var.set("Credenziali non corrette.")
            password_var.set("")
            password_entry.focus_set()
            password_entry.selection_range(0, "end")

        ttk.Button(button_row, text="Esci", command=dialog.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(button_row, text="Accedi", command=attempt_login).grid(row=0, column=1)

        dialog.bind("<Return>", attempt_login)
        dialog.bind("<Escape>", lambda event: dialog.destroy())
        username_entry.focus_set()

        dialog.update_idletasks()
        width = dialog.winfo_reqwidth()
        height = dialog.winfo_reqheight()
        screen_width = dialog.winfo_screenwidth()
        screen_height = dialog.winfo_screenheight()
        pos_x = max((screen_width - width) // 2, 0)
        pos_y = max((screen_height - height) // 2, 0)
        dialog.geometry(f"{width}x{height}+{pos_x}+{pos_y}")
        dialog.attributes("-topmost", True)
        dialog.lift()
        dialog.grab_set()

        self.wait_window(dialog)
        return authenticated["value"]

    def _build_variables(self) -> None:
        current_revision_datetime = current_datetime_display()
        self.vars = {
            "cliente": tk.StringVar(),
            "template_name": tk.StringVar(value=self.templates[0].name if self.templates else ""),
            "modello_dispositivo": tk.StringVar(),
            "norma": tk.StringVar(value="EN 365"),
            "codice_sap": tk.StringVar(),
            "lotto": tk.StringVar(),
            "serie": tk.StringVar(),
            "anno_fabbricazione": tk.StringVar(),
            "data_revisione": tk.StringVar(value=current_revision_datetime),
            "data_prossima_revisione": tk.StringVar(),
            "tecnico": tk.StringVar(),
            "esito": tk.StringVar(value="Positivo"),
            "export_pdf": tk.BooleanVar(value=True),
        }
        self.logo_var = tk.StringVar(value=self.settings.get("logo_path", ""))
        self.history_search_var = tk.StringVar()
        self.checklist_status_var = tk.StringVar(value="Scheda di controllo non caricata.")
        self.catalog_info_var = tk.StringVar(value="Catalogo Irudek: stato non disponibile.")
        self.status_var = tk.StringVar(value="Pronto.")

    def _load_settings(self) -> dict:
        if SETTINGS_PATH.exists():
            try:
                return json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
            except Exception:
                return default_settings()
        settings = default_settings()
        SETTINGS_PATH.write_text(json.dumps(settings, indent=2), encoding="utf-8")
        return settings

    def _save_settings(self) -> None:
        self.settings["logo_path"] = self.logo_var.get().strip()
        SETTINGS_PATH.write_text(json.dumps(self.settings, indent=2), encoding="utf-8")

    def _apply_window_icon(self) -> None:
        if not ICON_PATH.exists():
            return
        try:
            self.iconbitmap(default=str(ICON_PATH))
        except Exception:
            pass

    def _build_layout(self) -> None:
        self.columnconfigure(0, weight=3)
        self.columnconfigure(1, weight=2)
        self.rowconfigure(0, weight=1)

        form_frame = ttk.Frame(self, padding=14)
        form_frame.grid(row=0, column=0, sticky="nsew")
        form_frame.columnconfigure(1, weight=3)
        form_frame.columnconfigure(3, weight=2)

        history_frame = ttk.Frame(self, padding=(0, 14, 14, 14))
        history_frame.grid(row=0, column=1, sticky="nsew")
        history_frame.rowconfigure(2, weight=1)
        history_frame.columnconfigure(0, weight=1)

        ttk.Label(form_frame, text="Nuovo certificato", font=("Segoe UI", 15, "bold")).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(0, 10)
        )
        ttk.Label(form_frame, text=f"Utente: {self.current_user}", foreground="#425466").grid(
            row=0, column=2, sticky="e", padx=(10, 10), pady=(0, 10)
        )
        top_actions = ttk.Frame(form_frame)
        top_actions.grid(row=0, column=3, sticky="e", pady=(0, 10))
        self.catalog_button = ttk.Button(
            top_actions, text="Aggiorna catalogo Irudek", command=self._update_irudek_catalog
        )
        self.catalog_button.grid(row=0, column=0)
        ttk.Button(top_actions, text="Utenti", command=self._open_user_management).grid(row=0, column=1, padx=(8, 0))
        ttk.Label(top_actions, textvariable=self.catalog_info_var, foreground="#425466").grid(
            row=1, column=0, columnspan=2, sticky="e", pady=(6, 0)
        )

        ttk.Label(form_frame, text="Logo aziendale").grid(row=1, column=0, sticky="w", padx=(0, 8), pady=5)
        ttk.Entry(form_frame, textvariable=self.logo_var).grid(
            row=1, column=1, columnspan=2, sticky="ew", pady=5
        )
        ttk.Button(form_frame, text="Scegli logo", command=self._choose_logo).grid(
            row=1, column=3, sticky="ew", pady=5
        )

        ttk.Label(form_frame, text="Modello certificato").grid(
            row=2, column=0, sticky="w", padx=(0, 8), pady=5
        )
        template_widget = ttk.Combobox(
            form_frame,
            textvariable=self.vars["template_name"],
            values=list(self.template_lookup),
            state="readonly",
            width=76,
        )
        template_widget.bind("<<ComboboxSelected>>", self._sync_model_from_template)
        template_widget.grid(row=2, column=1, columnspan=3, sticky="ew", pady=5)

        fields = [
            ("Cliente", "cliente"),
            ("Modello dispositivo", "modello_dispositivo"),
            ("Norma", "norma"),
            ("Codice SAP", "codice_sap"),
            ("Numero lotto", "lotto"),
            ("Numero serie", "serie"),
            ("Anno fabbricazione", "anno_fabbricazione"),
            ("Data revisione", "data_revisione"),
            ("Data prossima revisione", "data_prossima_revisione"),
            ("Tecnico", "tecnico"),
            ("Esito", "esito"),
        ]

        row = 3
        for index, (label, key) in enumerate(fields):
            column = 0 if index % 2 == 0 else 2
            if index and index % 2 == 0:
                row += 1
            ttk.Label(form_frame, text=label).grid(row=row, column=column, sticky="w", padx=(0, 8), pady=5)
            if key == "cliente":
                widget = ttk.Combobox(form_frame, textvariable=self.vars[key])
                widget.bind("<<ComboboxSelected>>", self._load_client_defaults)
                widget.bind("<FocusOut>", self._load_client_defaults)
                self.client_widget = widget
            elif key == "modello_dispositivo":
                model_frame = ttk.Frame(form_frame)
                model_frame.columnconfigure(0, weight=1)
                widget = ttk.Entry(model_frame, textvariable=self.vars[key])
                widget.grid(row=0, column=0, sticky="ew")
                ttk.Button(model_frame, text="Trova Irudek", command=self._search_irudek_product_by_name).grid(
                    row=0, column=1, padx=(8, 0)
                )
                widget = model_frame
            elif key == "codice_sap":
                sap_frame = ttk.Frame(form_frame)
                sap_frame.columnconfigure(0, weight=1)
                sap_entry = ttk.Entry(sap_frame, textvariable=self.vars[key])
                sap_entry.grid(row=0, column=0, sticky="ew")
                sap_entry.bind("<FocusOut>", self._auto_fill_norma_from_sap)
                sap_entry.bind("<Return>", self._auto_fill_norma_from_sap)
                ttk.Button(sap_frame, text="Cerca SAP", command=self._auto_fill_norma_from_sap).grid(
                    row=0, column=1, padx=(8, 0)
                )
                widget = sap_frame
            elif key == "esito":
                widget = ttk.Combobox(
                    form_frame,
                    textvariable=self.vars[key],
                    values=["Positivo", "Negativo"],
                    state="readonly",
                )
            else:
                widget = ttk.Entry(form_frame, textvariable=self.vars[key])
            widget.grid(row=row, column=column + 1, sticky="ew", pady=5)

        row += 1
        ttk.Label(form_frame, text="Osservazioni").grid(row=row, column=0, sticky="nw", pady=(12, 5))
        self.osservazioni_widget = tk.Text(form_frame, height=7, wrap="word")
        self.osservazioni_widget.grid(row=row, column=1, columnspan=3, sticky="nsew", pady=(12, 5))
        form_frame.rowconfigure(row, weight=1)

        row += 1
        checklist_frame = ttk.Frame(form_frame)
        checklist_frame.grid(row=row, column=0, columnspan=4, sticky="ew", pady=(8, 4))
        checklist_frame.columnconfigure(1, weight=1)
        ttk.Label(checklist_frame, text="Scheda di controllo").grid(row=0, column=0, sticky="w")
        ttk.Label(checklist_frame, textvariable=self.checklist_status_var).grid(
            row=0, column=1, sticky="w", padx=(10, 0)
        )
        ttk.Button(checklist_frame, text="Compila scheda", command=self._open_checklist_editor).grid(
            row=0, column=2, padx=(12, 0)
        )

        row += 1
        controls = ttk.Frame(form_frame)
        controls.grid(row=row, column=0, columnspan=4, sticky="ew", pady=(12, 8))
        controls.columnconfigure(5, weight=1)

        ttk.Checkbutton(controls, text="Esporta anche PDF", variable=self.vars["export_pdf"]).grid(
            row=0, column=0, sticky="w"
        )
        self.generate_button = ttk.Button(controls, text="Genera certificato", command=self._generate_certificate)
        self.generate_button.grid(row=0, column=1, padx=(12, 0))
        self.generate_print_button = ttk.Button(
            controls, text="Genera e stampa 2 copie", command=self._generate_and_print
        )
        self.generate_print_button.grid(row=0, column=2, padx=(12, 0))
        self._refresh_generation_buttons()
        ttk.Button(controls, text="Pulisci campi", command=self._clear_form).grid(row=0, column=3, padx=(12, 0))
        ttk.Button(controls, text="Apri cartella output", command=self._open_output_folder).grid(
            row=0, column=4, padx=(12, 0)
        )

        ttk.Label(form_frame, textvariable=self.status_var, foreground="#1f4d37").grid(
            row=row + 1, column=0, columnspan=4, sticky="w", pady=(6, 0)
        )

        ttk.Label(history_frame, text="Storico revisioni", font=("Segoe UI", 14, "bold")).grid(
            row=0, column=0, sticky="w", pady=(0, 8)
        )
        search_frame = ttk.Frame(history_frame)
        search_frame.grid(row=1, column=0, sticky="ew", pady=(0, 8))
        search_frame.columnconfigure(1, weight=1)
        ttk.Label(search_frame, text="Ricerca cliente / lotto / seriale (usa &&)").grid(
            row=0, column=0, sticky="w", padx=(0, 8)
        )
        search_entry = ttk.Entry(search_frame, textvariable=self.history_search_var)
        search_entry.grid(row=0, column=1, sticky="ew")
        search_entry.bind("<KeyRelease>", self._on_history_search)

        columns = ("id", "created_at", "updated_at", "cliente", "modello", "seriale", "data_revisione", "esito")
        self.history = ttk.Treeview(history_frame, columns=columns, show="headings", height=24)
        headings = {
            "id": "ID",
            "created_at": "Creato il",
            "updated_at": "Ultima modifica",
            "cliente": "Cliente",
            "modello": "Modello",
            "seriale": "Seriale",
            "data_revisione": "Revisione",
            "esito": "Esito",
        }
        widths = {
            "id": 55,
            "created_at": 135,
            "updated_at": 135,
            "cliente": 150,
            "modello": 170,
            "seriale": 110,
            "data_revisione": 95,
            "esito": 80,
        }
        for key in columns:
            self.history.heading(key, text=headings[key])
            self.history.column(key, width=widths[key], minwidth=55, anchor="w", stretch=True)
        self.history.grid(row=2, column=0, sticky="nsew")
        self.history_column_weights = {
            "id": 1,
            "created_at": 2,
            "updated_at": 2,
            "cliente": 2,
            "modello": 3,
            "seriale": 2,
            "data_revisione": 2,
            "esito": 1,
        }
        self.bind("<Configure>", self._resize_history_columns)

        scrollbar = ttk.Scrollbar(history_frame, orient="vertical", command=self.history.yview)
        scrollbar.grid(row=2, column=1, sticky="ns")
        self.history.configure(yscrollcommand=scrollbar.set)

        history_actions = ttk.Frame(history_frame)
        history_actions.grid(row=3, column=0, sticky="ew", pady=(10, 0))
        ttk.Button(history_actions, text="Aggiorna", command=self._refresh_history).grid(row=0, column=0)
        ttk.Button(history_actions, text="Modifica selezionato", command=self._edit_selected_record).grid(
            row=0, column=1, padx=(10, 0)
        )
        ttk.Button(history_actions, text="Anteprima certificato", command=self._preview_selected_record).grid(
            row=0, column=2, padx=(10, 0)
        )
        ttk.Button(history_actions, text="Elimina selezionato", command=self._delete_selected_record).grid(
            row=0, column=3, padx=(10, 0)
        )

    def _choose_logo(self) -> None:
        selected = filedialog.askopenfilename(
            title="Seleziona logo aziendale",
            filetypes=[
                ("Immagini", "*.png *.jpg *.jpeg *.bmp *.gif"),
                ("Tutti i file", "*.*"),
            ],
        )
        if not selected:
            return
        self.logo_var.set(selected)
        self._save_settings()
        self.status_var.set(f"Logo aggiornato: {Path(selected).name}")

    def _open_user_management(self) -> None:
        if normalize_username(self.current_user) != DEFAULT_ADMIN_USERNAME:
            messagebox.showwarning("Accesso negato", "Solo l'utente admin puo' creare nuovi utenti.")
            return

        dialog = tk.Toplevel(self)
        dialog.title("Gestione utenti")
        dialog.resizable(False, False)
        dialog.transient(self)
        dialog.grab_set()

        if ICON_PATH.exists():
            try:
                dialog.iconbitmap(default=str(ICON_PATH))
            except Exception:
                pass

        frame = ttk.Frame(dialog, padding=18)
        frame.grid(row=0, column=0, sticky="nsew")
        frame.columnconfigure(1, weight=1)

        ttk.Label(frame, text="Crea nuovo utente", font=("Segoe UI", 12, "bold")).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(0, 10)
        )
        users_label = ttk.Label(frame, text="Utenti esistenti: " + ", ".join(self.user_store.list_usernames()))
        users_label.grid(row=1, column=0, columnspan=2, sticky="w", pady=(0, 12))

        username_var = tk.StringVar()
        password_var = tk.StringVar()
        confirm_var = tk.StringVar()

        ttk.Label(frame, text="Nuovo utente").grid(row=2, column=0, sticky="w", padx=(0, 10), pady=6)
        username_entry = ttk.Entry(frame, textvariable=username_var, width=28)
        username_entry.grid(row=2, column=1, sticky="ew", pady=6)
        ttk.Label(frame, text="Password").grid(row=3, column=0, sticky="w", padx=(0, 10), pady=6)
        password_entry = ttk.Entry(frame, textvariable=password_var, show="*", width=28)
        password_entry.grid(row=3, column=1, sticky="ew", pady=6)
        ttk.Label(frame, text="Conferma password").grid(row=4, column=0, sticky="w", padx=(0, 10), pady=6)
        confirm_entry = ttk.Entry(frame, textvariable=confirm_var, show="*", width=28)
        confirm_entry.grid(row=4, column=1, sticky="ew", pady=6)

        def save_user() -> None:
            if password_var.get() != confirm_var.get():
                messagebox.showerror("Password non valida", "Le password non coincidono.")
                return
            try:
                created_username = self.user_store.create_user(username_var.get(), password_var.get())
            except ValueError as exc:
                messagebox.showerror("Creazione utente", str(exc))
                return

            self.status_var.set(f"Nuovo utente creato: {created_username}")
            users_label.configure(text="Utenti esistenti: " + ", ".join(self.user_store.list_usernames()))
            username_var.set("")
            password_var.set("")
            confirm_var.set("")
            messagebox.showinfo("Utente creato", f"Utente salvato correttamente: {created_username}")
            dialog.destroy()

        button_row = ttk.Frame(frame)
        button_row.grid(row=5, column=0, columnspan=2, sticky="e", pady=(10, 0))
        ttk.Button(button_row, text="Chiudi", command=dialog.destroy).grid(row=0, column=0, padx=(0, 8))
        ttk.Button(button_row, text="Salva utente", command=save_user).grid(row=0, column=1)

        dialog.bind("<Return>", lambda event: save_user())
        dialog.bind("<Escape>", lambda event: dialog.destroy())

        dialog.update_idletasks()
        width = dialog.winfo_reqwidth()
        height = dialog.winfo_reqheight()
        screen_width = dialog.winfo_screenwidth()
        screen_height = dialog.winfo_screenheight()
        pos_x = max((screen_width - width) // 2, 0)
        pos_y = max((screen_height - height) // 2, 0)
        dialog.geometry(f"{width}x{height}+{pos_x}+{pos_y}")
        username_entry.focus_set()

    def _observations(self) -> str:
        if not self.osservazioni_widget:
            return ""
        return self.osservazioni_widget.get("1.0", "end").strip()

    def _record_from_form(self) -> RevisionRecord:
        return RevisionRecord(
            cliente=self.vars["cliente"].get().strip(),
            template_name=self.vars["template_name"].get().strip(),
            modello_dispositivo=self.vars["modello_dispositivo"].get().strip(),
            norma=self.vars["norma"].get().strip(),
            codice_sap=self.vars["codice_sap"].get().strip(),
            lotto=self.vars["lotto"].get().strip(),
            serie=self.vars["serie"].get().strip(),
            anno_fabbricazione=self.vars["anno_fabbricazione"].get().strip(),
            data_revisione=self.vars["data_revisione"].get().strip(),
            data_prossima_revisione=self.vars["data_prossima_revisione"].get().strip(),
            tecnico=self.vars["tecnico"].get().strip(),
            osservazioni=self._observations(),
            esito=self.vars["esito"].get().strip(),
        )

    def _default_checklist_value(self) -> str:
        return "Positivo" if self.vars["esito"].get().strip() != "Negativo" else "Negativo"

    def _load_checklist_for_current_template(self) -> None:
        template_name = self.vars["template_name"].get().strip()
        template_path = self.template_lookup.get(template_name)
        if not template_path:
            self.checklist_items = []
            self.checklist_vars = []
            self.checklist_status_var.set("Scheda di controllo non disponibile.")
            return

        self.checklist_items = extract_checklist_items(template_path)
        default_value = self._default_checklist_value()
        self.checklist_vars = [tk.StringVar(value=default_value) for _ in self.checklist_items]
        if self.checklist_items:
            self.checklist_status_var.set(
                f"{len(self.checklist_items)} controlli caricati. Default: {default_value}."
            )
        else:
            self.checklist_status_var.set("Scheda di controllo non disponibile.")

    def _count_checklist_values(self) -> tuple[int, int]:
        positives = sum(1 for var in self.checklist_vars if var.get() == "Positivo")
        negatives = sum(1 for var in self.checklist_vars if var.get() == "Negativo")
        return positives, negatives

    def _update_checklist_status(self) -> None:
        if not self.checklist_vars:
            self.checklist_status_var.set("Scheda di controllo non disponibile.")
            return
        positives, negatives = self._count_checklist_values()
        self.checklist_status_var.set(
            f"{len(self.checklist_vars)} controlli compilati. Positivi: {positives}, Negativi: {negatives}."
        )

    def _set_all_checklist_values(self, value: str) -> None:
        for var in self.checklist_vars:
            var.set(value)
        self._update_checklist_status()

    def _open_checklist_editor(self) -> None:
        if not self.checklist_items:
            self._load_checklist_for_current_template()
        if not self.checklist_items:
            messagebox.showinfo("Scheda di controllo", "Il modello selezionato non contiene una scheda compilabile.")
            return

        window = tk.Toplevel(self)
        window.title("Scheda di controllo")
        window.geometry("980x640")
        window.transient(self)
        window.grab_set()

        header = ttk.Frame(window, padding=12)
        header.pack(fill="x")
        ttk.Label(
            header,
            text=f"Controlli per {self.vars['modello_dispositivo'].get().strip() or self.vars['template_name'].get().strip()}",
            font=("Segoe UI", 12, "bold"),
        ).pack(anchor="w")
        ttk.Label(header, text="Segna Positivo o Negativo per ogni riga della scheda.").pack(anchor="w", pady=(4, 0))

        toolbar = ttk.Frame(window, padding=(12, 0, 12, 8))
        toolbar.pack(fill="x")
        ttk.Button(toolbar, text="Tutto positivo", command=lambda: self._set_all_checklist_values("Positivo")).pack(
            side="left"
        )
        ttk.Button(toolbar, text="Tutto negativo", command=lambda: self._set_all_checklist_values("Negativo")).pack(
            side="left", padx=(8, 0)
        )

        body = ttk.Frame(window, padding=(12, 0, 12, 12))
        body.pack(fill="both", expand=True)
        body.columnconfigure(0, weight=1)
        body.rowconfigure(0, weight=1)

        canvas = tk.Canvas(body, highlightthickness=0)
        scrollbar = ttk.Scrollbar(body, orient="vertical", command=canvas.yview)
        scroll_frame = ttk.Frame(canvas)
        scroll_frame.bind(
            "<Configure>",
            lambda event: canvas.configure(scrollregion=canvas.bbox("all")),
        )
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        for idx, (item, var) in enumerate(zip(self.checklist_items, self.checklist_vars), start=1):
            row_frame = ttk.Frame(scroll_frame, padding=(0, 6))
            row_frame.pack(fill="x", expand=True)
            ttk.Label(row_frame, text=f"{idx}. {item}", wraplength=700, justify="left").pack(side="left", fill="x", expand=True)
            ttk.Radiobutton(row_frame, text="Positivo", value="Positivo", variable=var).pack(side="left", padx=(10, 0))
            ttk.Radiobutton(row_frame, text="Negativo", value="Negativo", variable=var).pack(side="left", padx=(10, 0))

        footer = ttk.Frame(window, padding=12)
        footer.pack(fill="x")
        ttk.Button(footer, text="Chiudi", command=lambda: [self._update_checklist_status(), window.destroy()]).pack(
            side="right"
        )

    def _validate(self, record: RevisionRecord) -> list[str]:
        missing = []
        for label, value in [
            ("Modello certificato", record.template_name),
            ("Cliente", record.cliente),
            ("Modello dispositivo", record.modello_dispositivo),
            ("Data revisione", record.data_revisione),
            ("Tecnico", record.tecnico),
        ]:
            if not value:
                missing.append(label)
        return missing

    def _set_generation_busy(self, busy: bool) -> None:
        state = "disabled" if busy else "normal"
        if self.generate_button:
            self.generate_button.configure(state=state)
        if self.generate_print_button:
            self.generate_print_button.configure(state=state)

    def _refresh_generation_buttons(self) -> None:
        if self.generate_button:
            self.generate_button.configure(
                text="Salva modifiche" if self.editing_record_id else "Genera certificato"
            )
        if self.generate_print_button:
            self.generate_print_button.configure(
                text="Salva modifiche e stampa 2 copie"
                if self.editing_record_id
                else "Genera e stampa 2 copie"
            )

    def _write_error_log(self, message: str) -> None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        with LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {message}\n")

    def _refresh_catalog_info(self) -> None:
        generated = parse_catalog_datetime(self.norm_catalog.generated_at)
        if generated is None:
            self.catalog_info_var.set("Catalogo Irudek: non aggiornato.")
            return
        age_days = (datetime.now() - generated).days
        age_label = "aggiornato oggi" if age_days == 0 else f"aggiornato {age_days} giorni fa"
        self.catalog_info_var.set(
            f"Catalogo Irudek: {generated.strftime('%d/%m/%Y %H:%M')} | {age_label} | {len(self.norm_catalog.items)} codici"
        )

    def _set_catalog_busy(self, busy: bool) -> None:
        if self.catalog_button:
            self.catalog_button.configure(state="disabled" if busy else "normal")

    def _resize_history_columns(self, _event=None) -> None:
        if not hasattr(self, "history") or not self.history.winfo_exists():
            return
        total_width = self.history.winfo_width()
        if total_width <= 100:
            return
        columns = ("id", "created_at", "updated_at", "cliente", "modello", "seriale", "data_revisione", "esito")
        minimums = {
            "id": 50,
            "created_at": 110,
            "updated_at": 120,
            "cliente": 110,
            "modello": 140,
            "seriale": 90,
            "data_revisione": 95,
            "esito": 75,
        }
        available = max(total_width - 10, sum(minimums.values()))
        total_weight = sum(self.history_column_weights.values())
        allocated = 0
        for index, key in enumerate(columns):
            if index == len(columns) - 1:
                width = max(minimums[key], available - allocated)
            else:
                width = max(minimums[key], int(available * self.history_column_weights[key] / total_weight))
                allocated += width
            self.history.column(key, width=width, minwidth=minimums[key], stretch=True)

    def _update_irudek_catalog(self) -> None:
        self._set_catalog_busy(True)
        self.status_var.set("Aggiornamento catalogo Irudek in corso...")

        def worker() -> None:
            try:
                total_codes, conflicts = build_irudek_norm_catalog(IRUDEK_NORMS_PATH)
                self.norm_catalog.load()
                self.after(0, lambda: self._handle_catalog_update_success(total_codes, conflicts))
            except Exception as exc:
                error_text = f"{exc}\n\n{traceback.format_exc()}"
                self._write_error_log(error_text)
                self.after(0, lambda: self._handle_catalog_update_failure(str(exc)))

        threading.Thread(target=worker, daemon=True).start()

    def _handle_catalog_update_success(self, total_codes: int, conflicts: int) -> None:
        self._set_catalog_busy(False)
        self._refresh_catalog_info()
        self.status_var.set(f"Catalogo Irudek aggiornato. Codici disponibili: {total_codes}.")
        message = (
            "Catalogo Irudek aggiornato correttamente.\n\n"
            f"Codici SAP caricati: {total_codes}\n"
            f"Conflitti ignorati: {conflicts}\n\n"
            f"File aggiornato:\n{IRUDEK_NORMS_PATH}"
        )
        messagebox.showinfo("Catalogo aggiornato", message)

    def _handle_catalog_update_failure(self, error_message: str) -> None:
        self._set_catalog_busy(False)
        self._refresh_catalog_info()
        self.status_var.set("Aggiornamento catalogo Irudek non riuscito.")
        messagebox.showerror(
            "Errore aggiornamento catalogo",
            error_message + f"\n\nDettagli salvati in:\n{LOG_PATH}",
        )

    def _generate_certificate(self, print_after: bool = False) -> None:
        record = self._record_from_form()
        missing = self._validate(record)
        if missing:
            messagebox.showwarning("Campi mancanti", "Compila questi campi:\n- " + "\n- ".join(missing))
            return

        template_path = self.template_lookup.get(record.template_name)
        if not template_path:
            messagebox.showerror("Modello non trovato", "Il modello selezionato non e disponibile.")
            return

        logo_path = Path(self.logo_var.get().strip()) if self.logo_var.get().strip() else None
        if logo_path and not logo_path.exists():
            messagebox.showerror("Logo non trovato", "Il file del logo selezionato non esiste piu.")
            return

        self._set_generation_busy(True)
        editing_record_id = self.editing_record_id
        self.status_var.set("Aggiornamento revisione in corso..." if editing_record_id else "Generazione in corso...")

        def worker() -> None:
            try:
                docx_path, pdf_path, pdf_error = self.generator.generate(
                    template_path,
                    record,
                    export_pdf=self.vars["export_pdf"].get(),
                    logo_path=logo_path,
                    checklist_results=[var.get() for var in self.checklist_vars],
                )
                print_error = None
                if print_after:
                    try:
                        self.generator.print_document(docx_path, copies=2)
                    except Exception as exc:
                        print_error = str(exc)
                replaced_paths: list[Path] = []
                if editing_record_id:
                    existing_row = self.repository.get_record_by_id(editing_record_id)
                    if not existing_row:
                        raise RuntimeError("La revisione da modificare non esiste piu nello storico.")
                    self.repository.update_record(editing_record_id, record, docx_path, pdf_path)
                    for old_path in [existing_row["docx_path"], existing_row["pdf_path"]]:
                        if not old_path:
                            continue
                        old_file = Path(old_path)
                        if old_file not in [docx_path, pdf_path]:
                            replaced_paths.append(old_file)
                    record_id = editing_record_id
                else:
                    record_id = self.repository.insert_record(record, docx_path, pdf_path)
                self.after(
                    0,
                    lambda: self._handle_generation_success(
                        record_id, docx_path, pdf_path, pdf_error, print_error, print_after, bool(editing_record_id)
                    ),
                )
                for old_file in replaced_paths:
                    try:
                        if old_file.exists():
                            old_file.unlink()
                    except Exception as exc:
                        self._write_error_log(f"Cleanup warning for {old_file}: {exc}")
            except Exception as exc:
                error_text = f"{exc}\n\n{traceback.format_exc()}"
                self._write_error_log(error_text)
                self.after(0, lambda: self._handle_generation_failure(str(exc)))

        threading.Thread(target=worker, daemon=True).start()

    def _handle_generation_success(
        self,
        record_id: int,
        docx_path: Path,
        pdf_path: Path | None,
        pdf_error: str | None,
        print_error: str | None,
        print_after: bool,
        was_edit: bool,
    ) -> None:
        self._set_generation_busy(False)
        self._save_settings()
        self._refresh_client_values()
        self._refresh_history()
        self.editing_record_id = None
        self._refresh_generation_buttons()

        self.status_var.set(
            f"{'Revisione aggiornata' if was_edit else 'Certificato creato'}. ID storico {record_id}. File: {docx_path.name}"
            + (f" e {pdf_path.name}" if pdf_path else "")
            + (" | Stampa 2 copie inviata." if print_after and not print_error else "")
        )

        message = ("Revisione aggiornata correttamente.\n\n" if was_edit else "Certificato generato correttamente.\n\n") + f"DOCX: {docx_path}\n"
        if pdf_path:
            message += f"PDF: {pdf_path}\n"
        elif pdf_error:
            message += "PDF non creato, ma il Word e stato salvato.\n"
            message += f"Dettaglio PDF: {pdf_error}\n"
            self._write_error_log(f"PDF export warning for {docx_path.name}: {pdf_error}")
        if print_after:
            if print_error:
                message += f"Stampa non riuscita: {print_error}\n"
                self._write_error_log(f"Print warning for {docx_path.name}: {print_error}")
            else:
                message += "Stampa 2 copie inviata.\n"

        messagebox.showinfo("Operazione completata", message)

    def _handle_generation_failure(self, error_message: str) -> None:
        self._set_generation_busy(False)
        self.status_var.set("Generazione non riuscita.")
        messagebox.showerror(
            "Errore",
            error_message + f"\n\nDettagli salvati in:\n{LOG_PATH}",
        )

    def _generate_and_print(self) -> None:
        self._generate_certificate(print_after=True)

    def _clear_form(self) -> None:
        current_revision_datetime = current_datetime_display()
        self.editing_record_id = None
        self._refresh_generation_buttons()
        self.vars["cliente"].set("")
        self.vars["modello_dispositivo"].set("")
        self.vars["norma"].set("EN 365")
        self.vars["codice_sap"].set("")
        self.vars["lotto"].set("")
        self.vars["serie"].set("")
        self.vars["anno_fabbricazione"].set("")
        self.vars["data_revisione"].set(current_revision_datetime)
        self.vars["data_prossima_revisione"].set("")
        self.vars["tecnico"].set("")
        self.vars["esito"].set("Positivo")
        self.vars["export_pdf"].set(True)
        if self.osservazioni_widget:
            self.osservazioni_widget.delete("1.0", "end")
        self._sync_model_from_template()
        self.status_var.set("Campi ripuliti. Modalita modifica disattivata.")

    def _refresh_client_values(self) -> None:
        if self.client_widget:
            self.client_widget.configure(values=self.repository.distinct_clients())

    def _sync_model_from_template(self, _event=None) -> None:
        template_name = self.vars["template_name"].get().strip()
        if not template_name:
            return
        device_name = template_name.replace("CERTIFICATO DI REVISIONE - ", "").replace(".docx", "").strip()
        self.vars["modello_dispositivo"].set(device_name)
        self._load_checklist_for_current_template()

    def _load_client_defaults(self, _event=None) -> None:
        cliente = self.vars["cliente"].get().strip()
        if not cliente:
            return
        row = self.repository.latest_record_for_client(cliente)
        if not row:
            return

        template_name = row["template_name"] or ""
        if template_name in self.template_lookup:
            self.vars["template_name"].set(template_name)
            self._sync_model_from_template()

        self.vars["modello_dispositivo"].set(row["modello_dispositivo"] or self.vars["modello_dispositivo"].get())
        self.vars["norma"].set(row["norma"] or self.vars["norma"].get())
        self.vars["codice_sap"].set(row["codice_sap"] or "")
        self.vars["lotto"].set(row["lotto"] or "")
        self.vars["serie"].set(row["serie"] or "")
        self.vars["anno_fabbricazione"].set(row["anno_fabbricazione"] or "")
        self.vars["data_prossima_revisione"].set(row["data_prossima_revisione"] or "")
        self.vars["tecnico"].set(row["tecnico"] or self.vars["tecnico"].get())
        self.vars["esito"].set(row["esito"] or self.vars["esito"].get())
        if self.osservazioni_widget:
            self.osservazioni_widget.delete("1.0", "end")
            self.osservazioni_widget.insert("1.0", row["osservazioni"] or "")
        self.status_var.set(f"Dati recenti caricati per il cliente: {cliente}")

    def _auto_fill_norma_from_sap(self, _event=None) -> None:
        codice_sap = self.vars["codice_sap"].get().strip()
        if not codice_sap:
            return
        match = self.norm_catalog.lookup(codice_sap)
        if not match:
            return
        standard = str(match.get("standard", "")).strip()
        if not standard:
            return
        self.vars["norma"].set(standard)
        title = str(match.get("title", "")).strip()
        self.status_var.set(
            f"Norma caricata da catalogo Irudek per SAP {normalize_catalog_code(codice_sap)}"
            + (f": {title}" if title else "")
        )

    def _apply_irudek_catalog_match(self, code: str, entry: dict) -> None:
        normalized_code = normalize_catalog_code(code)
        self.vars["codice_sap"].set(normalized_code)
        standard = str(entry.get("standard", "")).strip()
        title = str(entry.get("title", "")).strip()
        if standard:
            self.vars["norma"].set(standard)
        if title:
            self.vars["modello_dispositivo"].set(title)
        self.status_var.set(
            f"Prodotto Irudek trovato: {title or normalized_code}"
            + (f" | Norma: {standard}" if standard else "")
        )

    def _choose_irudek_product_match(self, matches: list[tuple[str, dict]]) -> tuple[str, dict] | None:
        dialog = tk.Toplevel(self)
        dialog.title("Seleziona prodotto Irudek")
        dialog.geometry("760x360")
        dialog.transient(self)
        dialog.grab_set()

        frame = ttk.Frame(dialog, padding=14)
        frame.pack(fill="both", expand=True)
        ttk.Label(frame, text="Sono stati trovati piu prodotti. Seleziona quello corretto.", font=("Segoe UI", 11, "bold")).pack(
            anchor="w", pady=(0, 10)
        )

        columns = ("sap", "title", "standard")
        tree = ttk.Treeview(frame, columns=columns, show="headings", height=12)
        tree.heading("sap", text="Codice SAP")
        tree.heading("title", text="Prodotto")
        tree.heading("standard", text="Norma")
        tree.column("sap", width=120, anchor="w")
        tree.column("title", width=260, anchor="w")
        tree.column("standard", width=240, anchor="w")
        tree.pack(fill="both", expand=True)

        for code, entry in matches:
            tree.insert("", "end", iid=code, values=(code, entry.get("title", ""), entry.get("standard", "")))

        selected: dict[str, tuple[str, dict] | None] = {"value": None}

        def confirm_selection(event: object | None = None) -> None:
            picked = tree.selection()
            if not picked:
                return
            code = picked[0]
            entry = next((item[1] for item in matches if item[0] == code), None)
            if entry is None:
                return
            selected["value"] = (code, entry)
            dialog.destroy()

        buttons = ttk.Frame(frame)
        buttons.pack(fill="x", pady=(10, 0))
        ttk.Button(buttons, text="Annulla", command=dialog.destroy).pack(side="right")
        ttk.Button(buttons, text="Seleziona", command=confirm_selection).pack(side="right", padx=(0, 8))

        tree.bind("<Double-1>", confirm_selection)
        if matches:
            tree.selection_set(matches[0][0])
            tree.focus(matches[0][0])

        self.wait_window(dialog)
        return selected["value"]

    def _search_irudek_product_by_name(self) -> None:
        query = self.vars["modello_dispositivo"].get().strip()
        if not query:
            messagebox.showinfo("Ricerca Irudek", "Inserisci prima il nome del prodotto nel campo Modello dispositivo.")
            return
        matches = self.norm_catalog.search_by_name(query)
        if not matches:
            messagebox.showinfo("Ricerca Irudek", "Nessun prodotto Irudek trovato con questo nome.")
            return
        if len(matches) == 1:
            code, entry = matches[0]
            self._apply_irudek_catalog_match(code, entry)
            return
        choice = self._choose_irudek_product_match(matches)
        if not choice:
            return
        code, entry = choice
        self._apply_irudek_catalog_match(code, entry)

    def _on_history_search(self, _event=None) -> None:
        self._refresh_history()

    def _edit_selected_record(self) -> None:
        selected = self.history.selection()
        if not selected:
            messagebox.showwarning("Nessuna selezione", "Seleziona prima una riga nello storico.")
            return

        record_id = int(selected[0])
        row = self.repository.get_record_by_id(record_id)
        if not row:
            messagebox.showerror("Errore", "La riga selezionata non esiste piu nello storico.")
            self._refresh_history()
            return

        template_name = row["template_name"] or ""
        if template_name in self.template_lookup:
            self.vars["template_name"].set(template_name)
            self._sync_model_from_template()

        self.vars["cliente"].set(row["cliente"] or "")
        self.vars["modello_dispositivo"].set(row["modello_dispositivo"] or "")
        self.vars["norma"].set(row["norma"] or "")
        self.vars["codice_sap"].set(row["codice_sap"] or "")
        self.vars["lotto"].set(row["lotto"] or "")
        self.vars["serie"].set(row["serie"] or "")
        self.vars["anno_fabbricazione"].set(row["anno_fabbricazione"] or "")
        self.vars["data_revisione"].set(row["data_revisione"] or "")
        self.vars["data_prossima_revisione"].set(row["data_prossima_revisione"] or "")
        self.vars["tecnico"].set(row["tecnico"] or "")
        self.vars["esito"].set(row["esito"] or "Positivo")
        if self.osservazioni_widget:
            self.osservazioni_widget.delete("1.0", "end")
            self.osservazioni_widget.insert("1.0", row["osservazioni"] or "")
        self._load_checklist_for_current_template()
        self.editing_record_id = record_id
        self._refresh_generation_buttons()
        self.status_var.set(
            f"Modalita modifica attiva per ID storico {record_id}. Correggi i dati e premi Genera certificato per salvare."
        )

    def _preview_selected_record(self) -> None:
        selected = self.history.selection()
        if not selected:
            messagebox.showwarning("Nessuna selezione", "Seleziona prima una riga nello storico.")
            return

        record_id = int(selected[0])
        row = self.repository.get_record_by_id(record_id)
        if not row:
            messagebox.showerror("Errore", "La riga selezionata non esiste piu nello storico.")
            self._refresh_history()
            return

        candidates = []
        if row["pdf_path"]:
            candidates.append(Path(row["pdf_path"]))
        if row["docx_path"]:
            candidates.append(Path(row["docx_path"]))

        existing = next((path for path in candidates if path.exists()), None)
        if not existing:
            messagebox.showerror(
                "File non trovato",
                "Non riesco a trovare ne il PDF ne il Word collegato a questa revisione.",
            )
            return

        os.startfile(str(existing))
        self.status_var.set(f"Anteprima aperta: {existing.name}")

    def _delete_selected_record(self) -> None:
        selected = self.history.selection()
        if not selected:
            messagebox.showwarning("Nessuna selezione", "Seleziona prima una riga nello storico.")
            return

        record_id = int(selected[0])
        row = self.repository.get_record_by_id(record_id)
        if not row:
            messagebox.showerror("Errore", "La riga selezionata non esiste piu nello storico.")
            self._refresh_history()
            return

        docx_path = Path(row["docx_path"]) if row["docx_path"] else None
        pdf_path = Path(row["pdf_path"]) if row["pdf_path"] else None
        filename = docx_path.name if docx_path else f"record {record_id}"

        confirmed = messagebox.askyesno(
            "Conferma eliminazione",
            "Vuoi eliminare il file salvato e la riga nello storico?\n\n"
            f"Elemento: {filename}",
        )
        if not confirmed:
            return

        errors: list[str] = []
        for file_path in [docx_path, pdf_path]:
            if not file_path:
                continue
            try:
                if file_path.exists():
                    file_path.unlink()
            except Exception as exc:
                errors.append(f"{file_path.name}: {exc}")

        if errors:
            messagebox.showerror(
                "Eliminazione non completata",
                "Non sono riuscito a eliminare tutti i file:\n- " + "\n- ".join(errors),
            )
            return

        self.repository.delete_record(record_id)
        self._refresh_history()
        self._refresh_client_values()
        self.status_var.set(f"Elemento eliminato: {filename}")
        messagebox.showinfo("Eliminazione completata", f"Elemento eliminato correttamente:\n{filename}")

    def _refresh_history(self) -> None:
        for item in self.history.get_children():
            self.history.delete(item)
        for display_index, row in enumerate(self.repository.search_records(self.history_search_var.get()), start=1):
            self.history.insert(
                "",
                "end",
                iid=str(row["id"]),
                values=(
                    display_index,
                    row["created_at"],
                    row["updated_at"] or "",
                    row["cliente"],
                    row["modello_dispositivo"],
                    row["serie"],
                    row["data_revisione"],
                    row["esito"],
                ),
            )

    def _open_output_folder(self) -> None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        subprocess.run(["explorer", str(OUTPUT_DIR)], check=False)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    PDF_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    app = RevisionApp()
    if not app.is_authenticated:
        return
    app.mainloop()


if __name__ == "__main__":
    main()
